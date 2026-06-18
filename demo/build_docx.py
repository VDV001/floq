#!/usr/bin/env python3
"""Собирает demo/Floq_демо_сценарии.docx из скриншотов demo/screens/ + подписи."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
SCREENS = os.path.join(HERE, "screens")
WEB = os.path.join(HERE, "screens_web")  # пересжатые под вставку в docx
OUT = os.path.join(HERE, "Floq_демо_сценарии.docx")

# Скриншоты снимаются в 2x (2880px) — для docx это лишний вес и тормоза Word/Pages.
# Ужимаем до разумной ширины, сохраняя чёткость на печатном масштабе ~6.6".
MAX_W = 1400


def web_version(fname):
    """Вернуть путь к пересжатой версии (создав при необходимости)."""
    os.makedirs(WEB, exist_ok=True)
    src = os.path.join(SCREENS, fname)
    dst = os.path.join(WEB, fname)
    img = Image.open(src)
    if img.width > MAX_W:
        h = round(img.height * MAX_W / img.width)
        img = img.resize((MAX_W, h), Image.LANCZOS)
    img.save(dst, optimize=True)
    return dst

BLUE = RGBColor(0x00, 0x4A, 0xC6)
GREY = RGBColor(0x43, 0x46, 0x55)

# (файл, заголовок сценария, подпись — что происходит на экране)
SHOTS = [
    ("01-login.png", "Вход в систему",
     "Точка входа в Floq. Авторизация по email и паролю; ИИ-ассистент по продажам. "
     "Поддерживаются регистрация и вход через Google."),
    ("02-inbox.png", "Единый инбокс «Входящие»",
     "Все входящие лиды из Telegram и Email в одной ленте. Слева — этапы воронки "
     "(Новые, Квалифицированные, В диалоге, Фоллоуап, Закрытые) и фильтр по источнику, "
     "справа — ИИ-сводка: в системе 8 лидов, 2 ждут ответа."),
    ("03-lead-qualification.png", "Карточка лида с ИИ-квалификацией",
     "Алексей Смирнов (ООО «Ромашка»). ИИ автоматически определил потребность "
     "(внедрение CRM на 12 человек с интеграциями), оценил бюджет и сроки и выставил "
     "скоринг 82/100 с рекомендованным действием. Ниже — вся история переписки."),
    ("04-lead-draft.png", "ИИ-черновик ответа и подтверждение действия",
     "Марина Котова (Студия «Лайт»). Справа ИИ сгенерировал персональный черновик ответа — "
     "менеджеру остаётся отредактировать и отправить. Сверху — авто-предложение со ссылкой "
     "на запись на демо, ожидающее одобрения оператора (human-in-the-loop)."),
    ("05-pending-replies.png", "Очередь авто-ответов (Human-in-the-loop)",
     "Сообщения, которые система сгенерировала автоматически (например, ссылки на запись "
     "на демо при согласии клиента), и которые ждут одобрения оператора перед отправкой. "
     "Защита от ошибочных авто-отправок."),
    ("06-leads.png", "Раздел «Лиды»",
     "Лента потенциальных клиентов, написавших первыми. Здесь ИИ квалифицирует и "
     "оценивает каждый контакт; доступны фильтры «Непрочитанные» и «Приоритетные»."),
    ("07-pipeline.png", "Канбан-воронка продаж",
     "Визуальный путь лида по этапам: Новый → Квалифицирован → В диалоге → Фоллоуап → "
     "Закрыт. Сверху — метрики (конверсия, активные лиды) и Floq AI Инсайт, "
     "подсвечивающий сделки, требующие внимания."),
    ("08-automations.png", "Автоматизации",
     "Тумблеры авто-действий: ИИ-квалификация входящих, генерация черновиков ответов, "
     "авто-фоллоуапы по «молчащим» лидам, авто-отправка. Рутина выполняется без "
     "ручного труда."),
    ("09-prospects.png", "База проспектов для холодного аутрича",
     "Контакты с компанией, должностью, индустрией, статусом и результатом верификации "
     "email/телефона. Пополняется импортом из CSV, парсингом 2GIS и ручным добавлением."),
    ("10-sequences.png", "Секвенции — цепочки касаний",
     "Мультиканальные последовательности: Email → Telegram → Прозвон с настраиваемыми "
     "задержками между шагами. ИИ пишет текст под каждый шаг и канал."),
    ("11-outbound.png", "Очередь отправки",
     "Контроль качества ИИ-сообщений перед отправкой: счётчики отправлено/открыто/"
     "ответили/bounce, режим «Автопилот» и карточки сообщений с кнопкой подтверждения."),
    ("12-analytics-sequences.png", "Аналитика секвенций",
     "Сравнение эффективности последовательностей: sent / delivered / opened / replied / "
     "converted и проценты Open%, Reply%, Conv%. Сразу видно, какая цепочка работает лучше."),
    ("13-analytics-cost.png", "Аналитика затрат на ИИ",
     "Сколько потрачено на ИИ и что получено: общий расход и число вызовов, стоимость за "
     "лид / квалифицированного / конверсию / отправленное сообщение, разбивка по типу "
     "запроса и по модели (gpt-4o, claude-haiku, gpt-4o-mini, локальный gemma3)."),
    ("14-settings.png", "Настройки",
     "Подключение каналов (Telegram-бот, IMAP/Email, SMTP, Resend), выбор ИИ-провайдера "
     "и модели, уведомления и тестирование интеграций прямо из интерфейса."),
    ("15-plans.png", "Тарифные планы",
     "Три тарифа — Starter, Growth и Pro — с разными лимитами по лидам и набором "
     "возможностей (аутрич, парсинг, верификация, API, выделенный менеджер)."),
]


def main():
    doc = Document()
    # поля поуже, чтобы картинки были крупнее
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.7)
        s.top_margin = s.bottom_margin = Inches(0.7)

    # титул
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Floq")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Демонстрация работы сервиса — основные сценарии")
    rs.font.size = Pt(15)
    rs.font.color.rgb = GREY

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = intro.add_run(
        "AI-помощник для полного цикла B2B-продаж: входящие лиды (Telegram + Email), "
        "холодный аутрич и AI-квалификация с черновиками ответов. "
        "Ниже — 15 ключевых экранов на тестовых данных."
    )
    ri.font.size = Pt(11)
    ri.font.color.rgb = GREY

    doc.add_page_break()

    for i, (fname, title, caption) in enumerate(SHOTS, 1):
        path = web_version(fname) if os.path.exists(os.path.join(SCREENS, fname)) else os.path.join(SCREENS, fname)
        h = doc.add_paragraph()
        rh = h.add_run(f"{i}. {title}")
        rh.bold = True
        rh.font.size = Pt(16)
        rh.font.color.rgb = BLUE

        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(6.6))

        c = doc.add_paragraph()
        rc = c.add_run(caption)
        rc.font.size = Pt(11)
        rc.font.color.rgb = GREY

        if i < len(SHOTS):
            doc.add_page_break()

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
